import os
import pandas as pd
from docx import Document
from utils.functions import *

def procesar_planilla_fase_inicial(ruta_excel, hoja_seleccionada, carpeta_destino, mes, año, numero_armada):

    datos = pd.read_excel(ruta_excel, sheet_name=hoja_seleccionada)

    carpeta_principal = os.path.join(carpeta_destino, 'FASE INICIAL')
    os.makedirs(carpeta_principal, exist_ok=True)

    for i, fila in enumerate(datos.itertuples(index=False), start=1):
        docente = str(getattr(fila, "Docente", "N/A"))
        nombre_docente = limpiar_nombre_archivo(docente)
        carpeta_docente = os.path.join(carpeta_principal, nombre_docente)
        os.makedirs(carpeta_docente, exist_ok=True)

        ruc = limpiar_numero(getattr(fila, "N_Ruc", ""))
        descripcion_raw = str(getattr(fila, "Curso", ""))
        descripcion = redactar_cursos(descripcion_raw)
        cant_cursos = int(getattr(fila, "Cantidad_cursos", 0))
        disenio_cant_horas = cant_cursos * 4
        horas_disenio = f"{int(round(disenio_cant_horas))} horas de diseño de exámenes"
        clasif_valor = int(getattr(fila, "Examen_clasif", 0))
        direccion = str(getattr(fila, "Domicilio_docente", '')).strip()
        correo = str(getattr(fila, "Correo_personal", ''))
        celular = limpiar_numero(getattr(fila, "Numero_celular", ""))
        dni_docente = limpiar_numero(getattr(fila, "Numero_dni", ""))
        if len(dni_docente) < 8:
            dni_docente = dni_docente.zfill(8)

        categoria_valor = getattr(fila, "Categoria_monto", 1)
        if pd.isna(categoria_valor):
            categoria_valor = 1
        else:
            categoria_valor = int(categoria_valor)

        clasif_cant_horas = clasif_valor / categoria_valor
        if clasif_cant_horas == 1:
            horas_clasif = f"{int(round(clasif_cant_horas))} hora de examen de clasificación"
        else:
            horas_clasif = f"{int(round(clasif_cant_horas))} horas de examen de clasificación"

        if clasif_valor == 0:
            descripcion_final = f"{descripcion} y {horas_disenio}"
        else:
            descripcion_final = f"{descripcion}, {horas_disenio} y {horas_clasif}"

        monto_categoria = getattr(fila, "Categoria_monto", 0)
        monto_categoria_letras = monto_a_letras(monto_categoria)
        monto_total = getattr(fila, "Subtotal_pago", 0)
        monto_total_letras = monto_a_letras(monto_total)
        tipo_contrato = getattr(fila, "Estado_docente", "N/A")
        nro_contrato_val = getattr(fila, "Nro_Contrato", "N/A")
        try:
            nro_contrato = str(int(float(nro_contrato_val)))
        except (ValueError, TypeError):
            nro_contrato = str(nro_contrato_val)

        # -------- GENERAR OFICIO --------
        if tipo_contrato == "CONTRATO":
            plantilla_oficio = ruta_absoluta_relativa('./Modelos_documentos/oficio_contrato.docx')
        elif tipo_contrato == "TERCERO":
            plantilla_oficio = ruta_absoluta_relativa('./Modelos_documentos/oficio_tercero.docx')
        else:
            plantilla_oficio = None

        if plantilla_oficio and os.path.exists(plantilla_oficio):
            documento = Document(plantilla_oficio)
            reemplazos_oficio = {
                "Nro_Contrato": nro_contrato,
                "docente": docente,
                "descripcion": descripcion_final,
                "categoria": f"S/ {monto_categoria:,.2f} ({monto_categoria_letras})",
                "monto_subtotal": f"S/ {monto_total:,.2f} ({monto_total_letras})",
                "numero_armada": numero_armada
            }
            reemplazar_en_parrafos(documento, reemplazos_oficio)
            reemplazar_en_tablas(documento, reemplazos_oficio)
            ruta_salida_oficio = os.path.join(carpeta_docente, f"OFICIO - {nombre_docente} - {mes} {año}.docx")
            documento.save(ruta_salida_oficio)

        # -------- GENERAR TDR --------
        if tipo_contrato == "TERCERO":
            tipo_tdr = str(getattr(fila, "Categoria_letra", "")).strip().upper()
            plantilla_tdr = ruta_absoluta_relativa(f'./Modelos_documentos/tdr_tipo{tipo_tdr}_.docx')
            if os.path.exists(plantilla_tdr):
                documento_tdr = Document(plantilla_tdr)
                reemplazos_tdr = {
                    "descripcion": descripcion_final,
                    "categoria": f"S/ {monto_categoria:,.2f} ({monto_categoria_letras})",
                    "monto_subtotal": f"S/ {monto_total:,.2f} ({monto_total_letras})",
                }
                reemplazar_en_parrafos(documento_tdr, reemplazos_tdr)
                reemplazar_en_tablas(documento_tdr, reemplazos_tdr)
                ruta_salida_tdr = os.path.join(carpeta_docente, f"TDR - {nombre_docente} - {mes} {año}.docx")
                documento_tdr.save(ruta_salida_tdr)

        # -------- GENERAR COTIZACIÓN --------

        if tipo_contrato == "TERCERO":
            plantilla_cotizacion = ruta_absoluta_relativa('./Modelos_documentos/modelo_cotizacion.docx')
            if os.path.exists(plantilla_cotizacion):
                documento_cot = Document(plantilla_cotizacion)
                reemplazos = {
                    "nombre_docente": docente,
                    "direccion_cot": f"Dirección: {direccion}",
                    "ruc_docente_cot": f"RUC N.º {ruc}",
                    "correo_docente_cot": f"Correo: {correo}",
                    "celular_cot": f"Teléfono: {celular}",
                    "descripcion_servicio": descripcion_final,
                    "categoria_monto": f"S/ {monto_categoria:,.2f} ({monto_categoria_letras})",
                    "monto_subtotal": f"S/ {monto_total:,.2f} ({monto_total_letras})",
                    "dni_cot": f"DNI: {dni_docente}"
                }

                reemplazar_en_parrafos(documento_cot, reemplazos)
                reemplazar_en_tablas(documento_cot, reemplazos)

                ruta_salida_cot = os.path.join(carpeta_docente, f"COTIZACIÓN - {nombre_docente} - {mes} {año}.docx")
                documento_cot.save(ruta_salida_cot)

        print (f"{docente} - Documentos generados correctamente.")