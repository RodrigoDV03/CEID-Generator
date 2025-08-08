import os
import sys
import pandas as pd
from docx import Document
from .utils import *
from fuzzywuzzy import process
from fuzzywuzzy import fuzz

def ruta_absoluta_relativa(path_relativo):
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, path_relativo)

def generar_conformidad_desde_excel(fila, plantilla_path, ruta_salida, numero_armada):

    docente = str(getattr(fila, "Docente", "N/A"))
    nombre_docente = docente

    if not os.path.exists(plantilla_path):
        raise FileNotFoundError(f"No se encontró la plantilla: {plantilla_path}")

    ruc = limpiar_numero(getattr(fila, "N_Ruc", ""))
    descripcion_raw = str(getattr(fila, "Curso", ""))
    descripcion = redactar_cursos(descripcion_raw)
    cant_cursos = int(getattr(fila, "Cantidad_cursos", 0))
    disenio_cant_horas = cant_cursos * 4
    horas_disenio = f"{int(round(disenio_cant_horas))} horas de diseño de exámenes"
    clasif_valor = int(getattr(fila, "Examen_clasif", 0))

    monto_categoria = float(getattr(fila, "Categoria_monto", 1))
    if pd.isna(monto_categoria):
        monto_categoria = 1
    else:
        monto_categoria = int(monto_categoria)

    clasif_cant_horas = clasif_valor / monto_categoria if monto_categoria else 0

    if clasif_cant_horas == 1:
        horas_clasif = f"{int(round(clasif_cant_horas))} hora de examen de clasificación"
    else:
        horas_clasif = f"{int(round(clasif_cant_horas))} horas de examen de clasificación"

    if clasif_valor == 0:
        descripcion_final = f"{descripcion} y {horas_disenio}"
    else:
        descripcion_final = f"{descripcion}, {horas_disenio} y {horas_clasif}"

    monto_categoria_letras = monto_a_letras(monto_categoria)
    monto_total = getattr(fila, "Subtotal_pago", 0)
    try:
        monto_total = float(monto_total)
        monto_total_str = f"{monto_total:,.2f}"
    except Exception:
        monto_total_str = str(monto_total)
    monto_total_letras = monto_a_letras(monto_total)

    nro_contrato = getattr(fila, "Nro_Contrato", "")
    try:
        nro_contrato = int(float(nro_contrato))
    except (ValueError, TypeError):
        pass


    doc = Document(plantilla_path)
    for p in doc.paragraphs:
        for run in p.runs:
            if "nombre" in run.text:
                run.text = run.text.replace("nombre", str(nombre_docente))
            if "ruc" in run.text:
                run.text = run.text.replace("ruc", str(ruc))
            if "descripcion_cursos" in run.text:
                run.text = run.text.replace("descripcion_cursos", str(descripcion_final))
            if "monto_subtotal" in run.text:
                run.text = run.text.replace("monto_subtotal", f"S/. {monto_total_str} ({str(monto_total_letras)})")
            if "monto_hora" in run.text:
                run.text = run.text.replace("monto_hora", f"S/. {monto_categoria:.2f} ({str(monto_categoria_letras)})")
            if "Nro_Contrato" in run.text:
                run.text = run.text.replace("Nro_Contrato", str(nro_contrato))
            if "numero_armada" in run.text:
                run.text = run.text.replace("numero_armada", str(numero_armada))

    carpeta_final = os.path.dirname(ruta_salida)
    os.makedirs(carpeta_final, exist_ok=True)

    try:
        doc.save(ruta_salida)
    except Exception as e:
        print(f"Error al guardar {ruta_salida}: {e}")
        raise
    return ruta_salida


def procesar_planilla(ruta_excel, ruta_docente, hoja, carpeta_salida, mes, año, numero_armada):
    df = pd.read_excel(ruta_excel, sheet_name=hoja)
    df_control = pd.read_excel(ruta_docente)

    for _, fila in df.iterrows():
        try:
            estado = str(getattr(fila, "Contrato_o_tercero", "")).strip().upper()
            docente = str(getattr(fila, "Docente", "N/A"))
            nombre_docente = docente

            if estado == "CONTRATO":
                plantilla = ruta_absoluta_relativa("Modelos_documentos/CONFORMIDAD CONTRATO - MODELO.docx")
            elif estado == "TERCERO":
                plantilla = ruta_absoluta_relativa("Modelos_documentos/CONFORMIDAD TERCERO - MODELO.docx")
            else:
                raise ValueError(f"Estado inválido: {estado}")

            carpeta_final = os.path.join(carpeta_salida, "FASE FINAL", nombre_docente)
            os.makedirs(carpeta_final, exist_ok=True)

            # === GENERAR DOCUMENTO DE CONFORMIDAD ===
            nombre_archivo_conformidad = f"CONFORMIDAD - {nombre_docente} - {mes} {año}.docx"
            ruta_conformidad = os.path.join(carpeta_final, nombre_archivo_conformidad)

            generar_conformidad_desde_excel(fila, plantilla, ruta_conformidad, numero_armada)

            # === SI ES CONTRATO, GENERAR TAMBIÉN CONTROL DE AVANCE ===
            if estado == "CONTRATO":
                plantilla_control = ruta_absoluta_relativa("Modelos_documentos/Control de avance de pagos - MODELO.docx")
                nombre_archivo_control = f"CONTROL DE AVANCE - {nombre_docente} - {mes} {año}.docx"
                ruta_control = os.path.join(carpeta_final, nombre_archivo_control)

                try:
                    # Usa el monto_total calculado en generar_conformidad
                    monto_total = getattr(fila, "Subtotal_pago", 0)
                    monto_total = float(monto_total) if not pd.isna(monto_total) else 0

                    generar_control_avance(fila=fila, monto_subtotal=monto_total, ruta_salida=ruta_control, plantilla_control_path=plantilla_control, df_control=df_control)
                except Exception as e:
                    print(f"Error al generar control de avance para {nombre_docente}: {e}")
                    continue
        except Exception as e:
            print(f"Error procesando fila para {docente}: {e}")
            continue

        print(f"{docente} - Documentos generados correctamente.")


def generar_control_avance(fila, monto_subtotal, ruta_salida, plantilla_control_path, df_control):

    nombre_docente = str(getattr(fila, "Docente", ""))

    # Buscar coincidencia en Excel fijo
    resultado = process.extractOne(nombre_docente, df_control["Docente"], scorer=fuzz.token_sort_ratio)
    mejor_match, score = resultado[0], resultado[1]

    if score < 85:
        raise ValueError(f"No se encontró coincidencia adecuada para '{nombre_docente}' (score: {score})")

    fila_control = df_control[df_control["Docente"] == mejor_match].iloc[0]

    idioma_docente = str(fila_control["Especialidad"])
    monto_total = float(fila_control["Monto total"])
    nro_contrato = str(fila_control["Nro Contrato"])
    try:
        nro_contrato = int(float(nro_contrato))
    except (ValueError, TypeError):
        pass

    saldo_restante = monto_total - monto_subtotal

    # Cargar plantilla
    doc = Document(plantilla_control_path)

    reemplazos = {
        "Nombre_Docente": nombre_docente,
        "Idioma_Docente": idioma_docente,
        "Monto_Subtotal": f"S/ {monto_subtotal:,.2f}",
        "Monto_Total": f"S/ {monto_total:,.2f}",
        "Saldo_Restante": f"S/ {saldo_restante:,.2f}",
        "Nro_Contrato": str(nro_contrato)
    }

    for p in doc.paragraphs:
        for key, val in reemplazos.items():
            if key in p.text:
                p.text = p.text.replace(key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in reemplazos.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

    os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
    doc.save(ruta_salida)
    return ruta_salida