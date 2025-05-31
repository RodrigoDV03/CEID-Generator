import pandas as pd
from docx import Document
from num2words import num2words
import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox

def limpiar_nombre_archivo(nombre):
    return re.sub(r'[\\/*?:"<>|]', "", nombre)

def monto_a_letras(monto):
    try:
        monto = float(monto)
        entero = int(monto)
        centavos = int(round((monto - entero) * 100))
        return f"{num2words(entero, lang='es')} y {centavos:02d}/100 soles"
    except Exception:
        return "N/A"

def redactar_cursos(cadena):
    if not isinstance(cadena, str):
        return "N/A"
    cursos = [c.strip() for c in cadena.split("/") if c.strip()]
    if not cursos:
        return "N/A"
    resultado = f"servicio de dictado de 28 horas de clases de {cursos[0]}"
    for curso in cursos[1:]:
        resultado += f", 28 horas de clases de {curso}"
    return resultado

def generar_documentos(ruta_excel, hoja_seleccionada):
    datos = pd.read_excel(ruta_excel, sheet_name=hoja_seleccionada)

    carpeta_oficios = 'oficios_generados'
    carpeta_tdrs = 'tdrs_generados'
    os.makedirs(carpeta_oficios, exist_ok=True)
    os.makedirs(carpeta_tdrs, exist_ok=True)

    for i, fila in enumerate(datos.itertuples(index=False), start=1):        
        docente = str(getattr(fila, "Docente", "N/A"))
        descripcion_raw = str(getattr(fila, "Curso", ""))
        descripcion = redactar_cursos(descripcion_raw)
        cant_cursos = int(getattr(fila, "Cantidad_cursos", 0))
        disenio_cant_horas = cant_cursos * 4
        horas_disenio = f"{int(round(disenio_cant_horas))} horas de diseño de exámenes"
        clasif_valor = int(getattr(fila, "Examen_clasif", 0))

        categoria_valor = getattr(fila, "Categoria_monto", 1)
        if pd.isna(categoria_valor):
            categoria_valor = 1
        else:
            categoria_valor = int(categoria_valor)

        clasif_cant_horas = clasif_valor / categoria_valor
        horas_clasif = f"{int(round(clasif_cant_horas))} horas de clasificación"

        if clasif_valor == 0:
            descripcion_final = f"{descripcion} y {horas_disenio}"
        else:
            descripcion_final = f"{descripcion}, {horas_disenio} y {horas_clasif}"

        monto_categoria = getattr(fila, "Categoria_monto", 0)
        monto_categoria_letras = monto_a_letras(monto_categoria)
        monto_total = getattr(fila, "Subtotal_pago", 0)
        monto_total_letras = monto_a_letras(monto_total)

        nombre_docente = limpiar_nombre_archivo(docente)

        # -------- GENERAR OFICIO --------
        tipo_contrato = getattr(fila, "Contrato_o_tercero", "N/A")
        if tipo_contrato == "CONTRATO":
            plantilla_oficio = './Modelos_documentos/modelo_oficio_contrato_FLCH.docx'
        elif tipo_contrato == "TERCERO":
            plantilla_oficio = './Modelos_documentos/modelo_FLCH.docx'
        else:
            plantilla_oficio = None

        if plantilla_oficio and os.path.exists(plantilla_oficio):
            documento = Document(plantilla_oficio)
            for parrafo in documento.paragraphs:
                for run in parrafo.runs:
                    run.text = run.text.replace("docente", docente)
                    run.text = run.text.replace("descripcion", descripcion_final)
                    run.text = run.text.replace("categoria", f"S/ {monto_categoria:,.2f} ({monto_categoria_letras})")
                    run.text = run.text.replace("monto_subtotal", f"S/ {monto_total:,.2f} ({monto_total_letras})")
            ruta_salida_oficio = os.path.join(carpeta_oficios, f"OFICIO - {nombre_docente}.docx")
            documento.save(ruta_salida_oficio)

        # -------- GENERAR TDR --------
        tipo_tdr = str(getattr(fila, "Categoria_letra", "")).strip().upper()
        plantilla_tdr = f'./Modelos_documentos/tdr_tipo{tipo_tdr}_.docx'
        if os.path.exists(plantilla_tdr):
            documento_tdr = Document(plantilla_tdr)
            for parrafo in documento_tdr.paragraphs:
                for run in parrafo.runs:
                    run.text = run.text.replace("descripcion", descripcion_final)
                    run.text = run.text.replace("categoria", f"S/ {monto_categoria:,.2f} ({monto_categoria_letras})")
                    run.text = run.text.replace("monto_subtotal", f"S/ {monto_total:,.2f} ({monto_total_letras})")
            ruta_salida_tdr = os.path.join(carpeta_tdrs, f"TDR - {nombre_docente}.docx")
            documento_tdr.save(ruta_salida_tdr)

    messagebox.showinfo("Éxito", f"Oficios guardados en '{carpeta_oficios}'\nTDRs guardados en '{carpeta_tdrs}'")

# --- GUI ---
def seleccionar_archivo():
    ruta = filedialog.askopenfilename(
        title="Seleccionar archivo Excel",
        filetypes=[("Archivos Excel", "*.xlsx *.xls")]
    )
    if ruta:
        try:
            hojas = pd.ExcelFile(ruta).sheet_names
            hoja_var.set(hojas[0])  # por defecto
            hoja_menu['menu'].delete(0, 'end')
            for h in hojas:
                hoja_menu['menu'].add_command(label=h, command=lambda val=h: hoja_var.set(val))
            label_archivo.config(text=f"Archivo cargado: {os.path.basename(ruta)}", fg="green")
            boton_generar.config(state="normal")
            boton_generar.ruta_excel = ruta
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron leer las hojas: {e}")

def iniciar_generacion():
    hoja = hoja_var.get()
    ruta = boton_generar.ruta_excel
    if not hoja:
        messagebox.showerror("Error", "Debe seleccionar una hoja.")
        return
    generar_documentos(ruta, hoja)

root = tk.Tk()
root.title("Generador de Oficios y TDRs - CEID")
root.geometry("500x300")

label_archivo = tk.Label(root, text="Ningún archivo seleccionado", fg="red")
label_archivo.pack(pady=10)

tk.Button(root, text="Seleccionar archivo Excel", command=seleccionar_archivo).pack(pady=5)

hoja_var = tk.StringVar()
hoja_menu = tk.OptionMenu(root, hoja_var, "")
hoja_menu.pack(pady=10)

boton_generar = tk.Button(root, text="Generar documentos", state="disabled", command=iniciar_generacion)
boton_generar.pack(pady=20)

root.mainloop()