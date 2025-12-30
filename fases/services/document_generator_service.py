import os
from docx import Document
from typing import Dict, Optional
from fases.models import DocumentConfig


class DocumentGeneratorService:
    
    @staticmethod
    def generar_documento(
        ruta_template: str,
        reemplazos: Dict[str, str],
        ruta_salida: str,
        ruta_firma: Optional[str] = None
    ) -> bool:
        if not ruta_template or not os.path.exists(ruta_template):
            return False
        
        try:
            doc = Document(ruta_template)
            
            # Reemplazar texto (sin la firma)
            reemplazos_sin_firma = {
                k: v for k, v in reemplazos.items() 
                if k != "firma_docente"
            }
            
            DocumentGeneratorService._reemplazar_en_parrafos(doc, reemplazos_sin_firma)
            DocumentGeneratorService._reemplazar_en_tablas(doc, reemplazos_sin_firma)
            
            # Insertar firma si existe
            if ruta_firma and os.path.exists(ruta_firma):
                DocumentGeneratorService._insertar_firma(doc, ruta_firma)
            
            # Crear directorio si no existe
            os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
            
            doc.save(ruta_salida)
            return True
            
        except Exception as e:
            print(f"Error al generar documento: {e}")
            return False
    
    @staticmethod
    def _reemplazar_en_parrafos(documento: Document, reemplazos: Dict[str, str]):
        for parrafo in documento.paragraphs:
            for marcador, valor in reemplazos.items():
                if marcador in parrafo.text:
                    texto_nuevo = parrafo.text.replace(marcador, valor)
                    for run in parrafo.runs:
                        run.text = ''
                    if parrafo.runs:
                        parrafo.runs[0].text = texto_nuevo
    
    @staticmethod
    def _reemplazar_en_tablas(documento: Document, reemplazos: Dict[str, str]):
        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        for marcador, valor in reemplazos.items():
                            if marcador in parrafo.text:
                                texto_nuevo = parrafo.text.replace(marcador, valor)
                                for run in parrafo.runs:
                                    run.text = ''
                                if parrafo.runs:
                                    parrafo.runs[0].text = texto_nuevo
    
    @staticmethod
    def _insertar_firma(documento: Document, ruta_firma: str):
        # Reemplazar en párrafos
        for parrafo in documento.paragraphs:
            if "firma_docente" in parrafo.text:
                parrafo.text = ""
                run = parrafo.add_run()
                run.add_picture(ruta_firma)
        
        # Reemplazar en tablas
        for tabla in documento.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    for parrafo in celda.paragraphs:
                        if "firma_docente" in parrafo.text:
                            parrafo.text = ""
                            run = parrafo.add_run()
                            run.add_picture(ruta_firma)
    
    @staticmethod
    def eliminar_texto_administrativo(ruta_documento: str) -> bool:
        if not os.path.exists(ruta_documento):
            return False
        
        try:
            doc = Document(ruta_documento)
            texto_eliminar = [
                ", monto por hora: S/. 1.00 (uno y 00/100 soles)",
                "Monto por hora: S/. 1.00 (uno y 00/100 soles)"
            ]
            
            # Eliminar en párrafos
            for parrafo in doc.paragraphs:
                for run in parrafo.runs:
                    for texto in texto_eliminar:
                        if texto in run.text:
                            run.text = run.text.replace(texto, "")
            
            # Eliminar en tablas
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for parrafo in celda.paragraphs:
                            for run in parrafo.runs:
                                for texto in texto_eliminar:
                                    if texto in run.text:
                                        run.text = run.text.replace(texto, "")
            
            doc.save(ruta_documento)
            return True
            
        except Exception as e:
            print(f"Error al modificar documento administrativo: {e}")
            return False
    
    @staticmethod
    def obtener_ruta_plantilla(tipo_documento: str, tipo_contrato: str, config: DocumentConfig) -> str:
        from fases.functions import ruta_absoluta_relativa
        
        rutas_plantillas = {
            'oficio_contrato': './Modelos_documentos/oficio_contrato.docx',
            'oficio_tercero': './Modelos_documentos/oficio_tercero.docx',
            'tdr_administrativo': './Modelos_documentos/tdr_administrativo.docx',
            'cotizacion': './Modelos_documentos/modelo_cotizacion.docx',
            'conformidad_contrato': './Modelos_documentos/conformidad_contrato.docx',
            'conformidad_tercero': './Modelos_documentos/conformidad_tercero.docx',
            'control_primera': './Modelos_documentos/control_pagos_primera.docx',
            'control_segunda': './Modelos_documentos/control_pagos_segunda.docx',
            'control_tercera': './Modelos_documentos/control_pagos_tercera.docx'
        }
        
        # Determinar la clave según el contexto
        if tipo_documento == 'oficio':
            clave = f'oficio_{tipo_contrato.lower()}'
        elif tipo_documento == 'tdr' and config.es_administrativo:
            clave = 'tdr_administrativo'
        elif tipo_documento == 'conformidad':
            clave = f'conformidad_{tipo_contrato.lower()}'
        elif tipo_documento == 'control':
            clave = f'control_{config.numero_armada}'
        else:
            clave = tipo_documento
        
        return ruta_absoluta_relativa(rutas_plantillas.get(clave, ''))
