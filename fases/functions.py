import pandas as pd
import os
from docx import Document
from decimal import Decimal, ROUND_HALF_UP
from num2words import num2words
from fases.utils import TextUtils, PathUtils

# Mantener funciones antiguas por compatibilidad
def limpiar_nombre_archivo(nombre):
    return TextUtils.limpiar_nombre_archivo(nombre)


def limpiar_numero(valor):
    return TextUtils.limpiar_numero(valor)


def reemplazar_en_parrafos(documento, reemplazos):
    for parrafo in documento.paragraphs:
        for marcador, valor in reemplazos.items():
            if marcador in parrafo.text:
                texto_nuevo = parrafo.text.replace(marcador, valor)
                for run in parrafo.runs:
                    run.text = ''
                if parrafo.runs:
                    parrafo.runs[0].text = texto_nuevo

def reemplazar_en_tablas(documento, reemplazos):
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for marcador, valor in reemplazos.items():
                        if marcador in parrafo.text:
                            texto_nuevo = parrafo.text.replace(marcador, valor)
                            for run in parrafo.runs:
                                run.text = ''
                            if parrafo.runs:
                                parrafo.runs[0].text = texto_nuevo

def monto_a_letras(monto):
    try:
        monto = Decimal(str(monto)).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        entero = int(monto)
        centavos = int((monto - Decimal(entero)) * 100)
        return f"{num2words(entero, lang='es')} y {centavos:02d}/100 soles"
    except Exception:
        return "N/A"

def redactar_cursos(cadena, tiene_bono=False):
    if not isinstance(cadena, str):
        return "N/A"
    cursos = [c.strip() for c in cadena.split("/") if c.strip()]
    if not cursos:
        return "N/A"
    
    elementos_descripcion = []
    
    for curso in cursos:
        # Detectar si es el servicio de actualización de materiales
        if "Servicio de actualización de materiales de enseñanza" in curso:
            # Para el servicio de actualización, usar el texto tal como viene (sin agregar "28 horas de clases")
            elementos_descripcion.append(curso.strip())
        else:
            # Para cursos académicos normales, agregar el formato tradicional
            elementos_descripcion.append(f"28 horas de clases de {curso}")
    
    # Agregar el bono al final si existe
    if tiene_bono:
        elementos_descripcion.append("servicio de diseño y evaluación del examen anual")
    
    # Unir todos los elementos
    if len(elementos_descripcion) == 1:
        # Si solo hay un elemento y es el servicio de actualización o bono, no agregar nada más
        if "Servicio de actualización de materiales de enseñanza" in elementos_descripcion[0] or "servicio de diseño y evaluación del examen anual" in elementos_descripcion[0]:
            return elementos_descripcion[0]
        else:
            return f"servicio de dictado de {elementos_descripcion[0]}"
    elif len(elementos_descripcion) == 2:
        # Si hay servicio de actualización o bono como segundo elemento, usar coma
        if "Servicio de actualización de materiales de enseñanza" in elementos_descripcion[1] or "servicio de diseño y evaluación del examen anual" in elementos_descripcion[1]:
            return f"servicio de dictado de {elementos_descripcion[0]}, {elementos_descripcion[1].lower()}"
        else:
            return f"servicio de dictado de {elementos_descripcion[0]}, {elementos_descripcion[1]}"
    else:
        # Si hay más de 2 elementos, usar solo comas
        return f"servicio de dictado de {', '.join(elementos_descripcion)}"

def ruta_absoluta_relativa(path_relativo):
    return PathUtils.ruta_absoluta_relativa(path_relativo)

def generar_documento(modelo_path, reemplazos, ruta_salida, firma_path=None):
    if modelo_path and os.path.exists(modelo_path):
        doc = Document(modelo_path)

        # Reemplazar solo texto (sin la firma)
        reemplazos_sin_firma = {k: v for k, v in reemplazos.items() if k != "firma_docente"}
        reemplazar_en_parrafos(doc, reemplazos_sin_firma)
        reemplazar_en_tablas(doc, reemplazos_sin_firma)
        
        # Insertar la firma donde está el marcador
        if firma_path and os.path.exists(firma_path):
            for parrafo in doc.paragraphs:
                if "firma_docente" in parrafo.text:
                    parrafo.text = ""
                    run = parrafo.add_run()
                    run.add_picture(firma_path)

            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for parrafo in celda.paragraphs:
                            if "firma_docente" in parrafo.text:
                                parrafo.text = ""
                                run = parrafo.add_run()
                                run.add_picture(firma_path)

        doc.save(ruta_salida)
        return True
    return False

def formato_soles(valor):
    try:
        return f"S/ {float(valor):,.2f}"
    except Exception:
        return f"S/ {valor}"
    
def generar_monto_referencial(monto_sin_actualizacion, monto_sin_actualizacion_letras, servicio_actualizacion, servicio_actualizacion_letras, bono, bono_letras, monto_total, monto_total_letras):
    lineas = []
    
    # Agregar monto base si existe
    if monto_sin_actualizacion > 0:
        lineas.append(f"S/. {monto_sin_actualizacion:,.2f} ({monto_sin_actualizacion_letras}).")
    
    # Agregar servicio de actualización si existe
    if servicio_actualizacion > 0:
        lineas.append(f"S/. {servicio_actualizacion:,.2f} ({servicio_actualizacion_letras}) por servicio de actualización de materiales de enseñanza.")
    
    # Agregar bono si existe
    if bono > 0:
        lineas.append(f"S/. {bono:,.2f} ({bono_letras}) por servicio de diseño y evaluación del examen anual.")
    
    # Si hay más de un concepto, agregar el monto total
    if len(lineas) > 1:
        lineas.append(f"Monto total: S/. {monto_total:,.2f} ({monto_total_letras}). Incluye el impuesto y la contribución de ley")
    else:
        # Si solo hay un concepto, solo agregar la nota de impuestos
        lineas.append(f"Incluye el impuesto y la contribución de ley")
        lineas_texto = lineas[0]
        if not lineas_texto.endswith("."):
            lineas_texto += "."
        return lineas_texto.replace(".", ". Incluye el impuesto y la contribución de ley", 1)
    
    return "\n".join(lineas)