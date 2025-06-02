import pandas as pd
from docx import Document
from num2words import num2words
import os
import re
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime

def limpiar_nombre_archivo(nombre):
    return re.sub(r'[\\/*?:"<>|]', "", nombre)


def limpiar_numero(valor):
    return "" if pd.isna(valor) else str(valor).split('.')[0]


def reemplazar_en_parrafos(documento, reemplazos):
    for parrafo in documento.paragraphs:
        for marcador, valor in reemplazos.items():
            if marcador in parrafo.text:
                texto_nuevo = parrafo.text.replace(marcador, valor)
                for run in parrafo.runs:
                    run.text = ''
                if parrafo.runs:
                    parrafo.runs[0].text = texto_nuevo

def reemplazar_en_tablas(documento, reemplazos):
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    for marcador, valor in reemplazos.items():
                        if marcador in parrafo.text:
                            texto_nuevo = parrafo.text.replace(marcador, valor)
                            for run in parrafo.runs:
                                run.text = ''
                            if parrafo.runs:
                                parrafo.runs[0].text = texto_nuevo

def monto_a_letras(monto):
    try:
        monto = float(monto)
        entero = int(monto)
        centavos = int(round((monto - entero) * 100))
        return f"{num2words(entero, lang='es')} y {centavos:02d}/100 soles"
    except Exception:
        return "N/A"

def redactar_cursos(cadena):
    if not isinstance(cadena, str):
        return "N/A"
    cursos = [c.strip() for c in cadena.split("/") if c.strip()]
    if not cursos:
        return "N/A"
    resultado = f"servicio de dictado de 28 horas de clases de {cursos[0]}"
    for curso in cursos[1:]:
        resultado += f", 28 horas de clases de {curso}"
    return resultado

def generar_documentos(ruta_excel, hoja_seleccionada):

    meses = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]

    hoy = datetime.now()
    fecha_actual = f"{hoy.day} de {meses[hoy.month - 1]} de {hoy.year}"

    datos = pd.read_excel(ruta_excel, sheet_name=hoja_seleccionada)

    carpeta_principal = 'FASE INICIAL'
    os.makedirs(carpeta_principal, exist_ok=True)

    for i, fila in enumerate(datos.itertuples(index=False), start=1):
        docente = str(getattr(fila, "Docente", "N/A"))
        nombre_docente = limpiar_nombre_archivo(docente)
        carpeta_docente = os.path.join(carpeta_principal, nombre_docente)
        os.makedirs(carpeta_docente, exist_ok=True)

        ruc = limpiar_numero(getattr(fila, "N_Ruc", ""))
        descripcion_raw = str(getattr(fila, "Curso", ""))
        descripcion = redactar_cursos(descripcion_raw)
        cant_cursos = int(getattr(fila, "Cantidad_cursos", 0))
        disenio_cant_horas = cant_cursos * 4
        horas_disenio = f"{int(round(disenio_cant_horas))} horas de diseño de exámenes"
        clasif_valor = int(getattr(fila, "Examen_clasif", 0))
        direccion = str(getattr(fila, "Domicilio_docente", '')).strip()
        correo = str(getattr(fila, "Correo_personal", ''))
        celular = limpiar_numero(getattr(fila, "Numero_celular", ""))
        dni_docente = limpiar_numero(getattr(fila, "Numero_dni", ""))
        if len(dni_docente) < 8:
            dni_docente = dni_docente.zfill(8)  # Añade ceros a la izquierda

        categoria_valor = getattr(fila, "Categoria_monto", 1)
        if pd.isna(categoria_valor):
            categoria_valor = 1
        else:
            categoria_valor = int(categoria_valor)

        clasif_cant_horas = clasif_valor / categoria_valor
        horas_clasif = f"{int(round(clasif_cant_horas))} horas de clasificación"

        if clasif_valor == 0:
            descripcion_final = f"{descripcion} y {horas_disenio}"
        else:
            descripcion_final = f"{descripcion}, {horas_disenio} y {horas_clasif}"

        monto_categoria = getattr(fila, "Categoria_monto", 0)
        monto_categoria_letras = monto_a_letras(monto_categoria)
        monto_total = getattr(fila, "Subtotal_pago", 0)
        monto_total_letras = monto_a_letras(monto_total)

        # -------- GENERAR OFICIO --------
        tipo_contrato = getattr(fila, "Contrato_o_tercero", "N/A")
        if tipo_contrato == "CONTRATO":
            plantilla_oficio = './Modelos_documentos/modelo_oficio_contrato_FLCH.docx'
        elif tipo_contrato == "TERCERO":
            plantilla_oficio = './Modelos_documentos/modelo_FLCH.docx'
        else:
            plantilla_oficio = None

        if plantilla_oficio and os.path.exists(plantilla_oficio):
            documento = Document(plantilla_oficio)
            for parrafo in documento.paragraphs:
                for run in parrafo.runs:
                    run.text = run.text.replace("docente", docente)
                    run.text = run.text.replace("descripcion", descripcion_final)
                    run.text = run.text.replace("categoria", f"S/ {monto_categoria:,.2f} ({monto_categoria_letras})")
                    run.text = run.text.replace("monto_subtotal", f"S/ {monto_total:,.2f} ({monto_total_letras})")
            ruta_salida_oficio = os.path.join(carpeta_docente, f"OFICIO - {nombre_docente}.docx")
            documento.save(ruta_salida_oficio)

        # -------- GENERAR TDR --------
        tipo_tdr = str(getattr(fila, "Categoria_letra", "")).strip().upper()
        plantilla_tdr = f'./Modelos_documentos/tdr_tipo{tipo_tdr}_.docx'
        if os.path.exists(plantilla_tdr):
            documento_tdr = Document(plantilla_tdr)
            for parrafo in documento_tdr.paragraphs:
                for run in parrafo.runs:
                    run.text = run.text.replace("descripcion", descripcion_final)
                    run.text = run.text.replace("categoria", f"S/ {monto_categoria:,.2f} ({monto_categoria_letras})")
                    run.text = run.text.replace("monto_subtotal", f"S/ {monto_total:,.2f} ({monto_total_letras})")
            ruta_salida_tdr = os.path.join(carpeta_docente, f"TDR - {nombre_docente}.docx")
            documento_tdr.save(ruta_salida_tdr)

        # -------- GENERAR COTIZACIÓN --------

        plantilla_cotizacion = './Modelos_documentos/modelo_cotizacion.docx'
        if os.path.exists(plantilla_cotizacion):
            documento_cot = Document(plantilla_cotizacion)
            reemplazos = {
                "nombre_docente": docente,
                "direccion_cot": f"Dirección: {direccion}",
                "ruc_docente_cot": f"RUC N.º {ruc}",
                "correo_docente_cot": f"Correo: {correo}",
                "celular_cot": f"Teléfono: {celular}",
                "fecha": fecha_actual,
                "descripcion_servicio": descripcion_final,
                "categoria_monto": f"S/ {monto_categoria:,.2f} ({monto_categoria_letras})",
                "monto_subtotal": f"S/ {monto_total:,.2f} ({monto_total_letras})",
                "dni_cot": f"DNI: {dni_docente}"
            }

            reemplazar_en_parrafos(documento_cot, reemplazos)
            reemplazar_en_tablas(documento_cot, reemplazos)

            ruta_salida_cot = os.path.join(carpeta_docente, f"COTIZACIÓN - {nombre_docente}.docx")
            documento_cot.save(ruta_salida_cot)

    messagebox.showinfo("Éxito", f"Todos los documentos se guardaron en la carpeta '{carpeta_principal}'.")

# --- VENTANA PRINCIPAL ---
root = tk.Tk()
root.title("Generador de Archivos Fase Inicial CEID")
root.geometry("560x380")
root.configure(bg="#f4f6fa")
root.resizable(False, False)

# --- ESTILOS ---
PRIMARY_COLOR = "#2d415a"
ACCENT_COLOR = "#4a90e2"
BG_COLOR = "#f4f6fa"
TEXT_COLOR = "#333"
DISABLED_COLOR = "#bbb"

FONT_TITLE = ("Segoe UI", 17, "bold")
FONT_LABEL = ("Segoe UI", 11)
FONT_BUTTON = ("Segoe UI", 10)
FONT_FOOTER = ("Segoe UI", 9)

# --- ENCABEZADO ---
tk.Label(
    root, text="Generador de Archivos Fase Inicial CEID",
    font=FONT_TITLE, bg=BG_COLOR, fg=PRIMARY_COLOR
).pack(pady=(20, 10))

ttk.Separator(root, orient="horizontal").pack(fill="x", padx=30)

# --- SELECCIÓN DE ARCHIVO ---
frame_archivo = tk.Frame(root, bg=BG_COLOR)
frame_archivo.pack(fill="x", padx=30, pady=(20, 5))

label_archivo = tk.Label(
    frame_archivo,
    text="Ningún archivo seleccionado",
    fg="red", bg=BG_COLOR,
    font=("Segoe UI", 10, "italic")
)
label_archivo.pack(side="left", padx=(0, 10))

def seleccionar_archivo():
    ruta = filedialog.askopenfilename(
        title="Seleccionar archivo Excel",
        filetypes=[("Archivos Excel", "*.xlsx *.xls")]
    )
    if ruta:
        try:
            hojas = pd.ExcelFile(ruta).sheet_names
            hoja_var.set(hojas[0])
            hoja_menu['menu'].delete(0, 'end')
            for h in hojas:
                hoja_menu['menu'].add_command(label=h, command=lambda val=h: hoja_var.set(val))
            label_archivo.config(text=f"📁 {os.path.basename(ruta)}", fg="green")
            boton_generar.config(state="normal")
            boton_generar.ruta_excel = ruta
        except Exception as e:
            messagebox.showerror("Error", f"No se pudieron leer las hojas:\n{e}")

tk.Button(
    frame_archivo,
    text="Seleccionar archivo",
    command=seleccionar_archivo,
    font=FONT_BUTTON,
    bg=ACCENT_COLOR, fg="white",
    activebackground="#357ABD", activeforeground="white",
    relief="flat", padx=10, pady=4
).pack(side="right")

# --- SELECCIÓN DE HOJA ---
frame_hoja = tk.Frame(root, bg=BG_COLOR)
frame_hoja.pack(fill="x", padx=30, pady=(15, 0))

tk.Label(
    frame_hoja,
    text="Selecciona la hoja de trabajo:",
    font=FONT_LABEL,
    bg=BG_COLOR, fg=TEXT_COLOR
).pack(side="left")

hoja_var = tk.StringVar()
hoja_menu = tk.OptionMenu(frame_hoja, hoja_var, "")
hoja_menu.config(font=FONT_BUTTON, width=25)
hoja_menu.pack(side="left", padx=(10, 0))

# --- BOTÓN GENERAR ---
def iniciar_generacion():
    hoja = hoja_var.get()
    ruta = getattr(boton_generar, "ruta_excel", None)
    if not hoja or not ruta:
        messagebox.showerror("Error", "Debe seleccionar un archivo y una hoja.")
        return
    generar_documentos(ruta, hoja)

boton_generar = tk.Button(
    root,
    text="📄 Generar documentos",
    state="disabled",
    command=iniciar_generacion,
    font=("Segoe UI", 12, "bold"),
    bg=PRIMARY_COLOR, fg="white",
    activebackground="#466a8f",
    activeforeground="white",
    relief="flat",
    padx=12, pady=8
)
boton_generar.pack(pady=35)

# --- PIE DE PÁGINA ---
tk.Label(
    root,
    text="CEID Generator - v1.0",
    font=FONT_FOOTER,
    bg=BG_COLOR,
    fg="#a0a8b8"
).pack(side="bottom", pady=(0, 10))

root.mainloop()